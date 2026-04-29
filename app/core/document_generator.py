"""Generador de documentos Word y texto plano. DRY con métodos privados reutilizables."""
import json
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path

from app.utils.logger import get_logger

logger = get_logger()

# Colores corporativos InnoTech (hex sin #, para docx RGBColor)
_COLOR_ACCENT   = (0xF9, 0x73, 0x16)   # naranja
_COLOR_DARK     = (0x0D, 0x1B, 0x2A)   # azul oscuro
_COLOR_MUTED    = (0x64, 0x74, 0x8B)   # gris texto secundario
_FONT_NAME      = "Calibri"


@dataclass
class GuionData:
    titulo: str
    funcionalidad: str
    palabras_clave: list[str]
    introduccion: str
    pasos: list[str]
    cierre: str
    resumen: str


@dataclass
class ProcessingMetadata:
    video_name: str
    video_path: str
    duration_seconds: float
    language_detected: str
    transcription_confidence: float
    ai_provider: str
    ai_model: str
    processed_at: str
    output_dir: str


class DocumentGenerator:
    """Genera archivos de salida: .docx, .txt, .json."""

    # ─── Públicos ─────────────────────────────────────────────────────────────

    def generate_transcription_txt(self, result, path: Path) -> None:
        path.write_text(result.to_timestamped_text(), encoding="utf-8")
        logger.info("Transcripción TXT guardada: %s", path)

    def generate_transcription_docx(self, result, path: Path) -> None:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        self._configure_margins(doc)
        self._add_header(doc, "Transcripción del Video")
        self._add_meta_line(doc, f"Confianza promedio: {result.get_word_confidence_avg():.1%}")

        # Tabla de segmentos con timestamps
        self._add_section_heading(doc, "Desglose por Segmentos")
        headers = ["Inicio", "Fin", "Segmento transcrito"]
        rows = [
            [self._fmt_time(s.start), self._fmt_time(s.end), s.text.strip()]
            for s in result.segments
        ]
        self._add_table(doc, headers, rows)

        # Texto completo
        self._add_section_heading(doc, "Transcripción Completa")
        self._add_body(doc, result.full_text)

        doc.save(str(path))
        logger.info("Transcripción DOCX guardada: %s", path)

    def generate_resumen_docx(self, resumen: str, titulo: str,
                              palabras_clave: list[str], path: Path) -> None:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        self._configure_margins(doc)
        self._add_header(doc, "Resumen Ejecutivo")
        self._add_meta_line(doc, titulo)

        self._add_section_heading(doc, "Resumen")
        self._add_body(doc, resumen)

        if palabras_clave:
            self._add_section_heading(doc, "Palabras Clave")
            self._add_body(doc, "  •  ".join(palabras_clave))

        doc.save(str(path))
        logger.info("Resumen DOCX guardado: %s", path)

    def generate_guion_base_docx(self, data: GuionData, path: Path) -> None:
        from docx import Document

        doc = Document()
        self._configure_margins(doc)
        self._add_header(doc, "Guión Base del Tutorial")
        self._add_meta_line(doc, data.titulo)

        if data.funcionalidad:
            self._add_section_heading(doc, "Funcionalidad")
            self._add_body(doc, data.funcionalidad)

        if data.palabras_clave:
            self._add_section_heading(doc, "Palabras Clave")
            self._add_body(doc, "  •  ".join(data.palabras_clave))

        if data.resumen:
            self._add_section_heading(doc, "Resumen Ejecutivo")
            self._add_body(doc, data.resumen)

        self._add_section_heading(doc, "Guión Base")

        if data.introduccion:
            self._add_subsection_heading(doc, "Introducción")
            self._add_body(doc, data.introduccion)

        if data.pasos:
            self._add_subsection_heading(doc, "Pasos del Tutorial")
            for i, paso in enumerate(data.pasos, 1):
                self._add_numbered_item(doc, i, paso)

        if data.cierre:
            self._add_subsection_heading(doc, "Cierre")
            self._add_body(doc, data.cierre)

        doc.save(str(path))
        logger.info("Guión base DOCX guardado: %s", path)

    def generate_complete_docx(self, transcription, guion_data: "GuionData", path: Path) -> None:
        """Genera un único DOCX con resumen, guión base y transcripción."""
        from docx import Document

        doc = Document()
        self._configure_margins(doc)
        self._add_header(doc, guion_data.titulo or "Informe de Videotutorial")
        self._add_meta_line(
            doc,
            f"InnoTech VideoTutor  ·  {datetime.now().strftime('%d/%m/%Y %H:%M')}",
        )
        if guion_data.palabras_clave:
            self._add_meta_line(doc, "  ·  ".join(guion_data.palabras_clave))

        # ── 1. Resumen ejecutivo ──────────────────────────────────────────────
        self._add_section_heading(doc, "1. Resumen Ejecutivo")
        self._add_body(doc, guion_data.resumen)

        # ── 2. Guión Base ─────────────────────────────────────────────────────
        self._add_section_heading(doc, "2. Guión Base")

        if guion_data.funcionalidad:
            self._add_subsection_heading(doc, "Funcionalidad")
            self._add_body(doc, guion_data.funcionalidad)

        if guion_data.introduccion:
            self._add_subsection_heading(doc, "Introducción")
            self._add_body(doc, guion_data.introduccion)

        if guion_data.pasos:
            self._add_subsection_heading(doc, "Pasos del Tutorial")
            for i, paso in enumerate(guion_data.pasos, 1):
                self._add_numbered_item(doc, i, paso)

        if guion_data.cierre:
            self._add_subsection_heading(doc, "Cierre")
            self._add_body(doc, guion_data.cierre)

        # ── 3. Transcripción ──────────────────────────────────────────────────
        self._add_section_heading(doc, "3. Transcripción Completa")
        self._add_meta_line(
            doc, f"Confianza promedio: {transcription.get_word_confidence_avg():.1%}"
        )

        if transcription.segments:
            self._add_subsection_heading(doc, "Desglose por Segmentos")
            self._add_table(
                doc,
                ["Inicio", "Fin", "Segmento"],
                [
                    [self._fmt_time(s.start), self._fmt_time(s.end), s.text.strip()]
                    for s in transcription.segments
                ],
            )

        self._add_subsection_heading(doc, "Texto Completo")
        self._add_body(doc, transcription.full_text)

        doc.save(str(path))
        logger.info("Informe completo DOCX guardado: %s", path)

    def generate_metadata_json(self, data: ProcessingMetadata, path: Path) -> None:
        path.write_text(
            json.dumps(asdict(data), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        logger.info("Metadata JSON guardada: %s", path)

    # ─── Privados: estructura ─────────────────────────────────────────────────

    def _configure_margins(self, doc) -> None:
        from docx.shared import Inches
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

    def _add_header(self, doc, title: str) -> None:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Línea de marca
        brand = doc.add_paragraph()
        brand.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = brand.add_run("InnoTech Solutions  |  VideoTutor")
        run.font.name = _FONT_NAME
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*_COLOR_MUTED)
        run.font.bold = False

        # Separador visual (regla)
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run("─" * 80)
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(*_COLOR_MUTED)

        # Título principal
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after  = Pt(4)
        run = heading.add_run(title)
        run.font.name  = _FONT_NAME
        run.font.size  = Pt(22)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(*_COLOR_DARK)

    def _add_meta_line(self, doc, text: str) -> None:
        from docx.shared import Pt, RGBColor
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.font.name  = _FONT_NAME
        run.font.size  = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(*_COLOR_MUTED)

    def _add_section_heading(self, doc, text: str) -> None:
        from docx.shared import Pt, RGBColor
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text.upper())
        run.font.name  = _FONT_NAME
        run.font.size  = Pt(11)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(*_COLOR_ACCENT)

    def _add_subsection_heading(self, doc, text: str) -> None:
        from docx.shared import Pt, RGBColor
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(text)
        run.font.name  = _FONT_NAME
        run.font.size  = Pt(11)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(*_COLOR_DARK)

    def _add_body(self, doc, text: str) -> None:
        from docx.shared import Pt, RGBColor
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.name  = _FONT_NAME
        run.font.size  = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

    def _add_numbered_item(self, doc, number: int, text: str) -> None:
        from docx.shared import Pt, RGBColor, Inches
        p = doc.add_paragraph()
        p.paragraph_format.space_after      = Pt(5)
        p.paragraph_format.left_indent      = Inches(0.2)
        # Número en naranja
        num_run = p.add_run(f"{number}. ")
        num_run.font.name  = _FONT_NAME
        num_run.font.size  = Pt(11)
        num_run.font.bold  = True
        num_run.font.color.rgb = RGBColor(*_COLOR_ACCENT)
        # Texto del paso
        txt_run = p.add_run(text)
        txt_run.font.name  = _FONT_NAME
        txt_run.font.size  = Pt(11)
        txt_run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

    def _add_table(self, doc, headers: list[str], rows: list[list[str]]) -> None:
        from docx.shared import Pt, RGBColor

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        # Cabecera con fondo naranja
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(header)
            run.font.name  = _FONT_NAME
            run.font.size  = Pt(10)
            run.font.bold  = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            self._set_cell_bg(cell, "F97316")

        # Filas de datos
        for row_data in rows:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                cell = row.cells[i]
                cell.text = ""
                run = cell.paragraphs[0].add_run(cell_text)
                run.font.name = _FONT_NAME
                run.font.size = Pt(10)

        doc.add_paragraph()  # espacio tras tabla

    def _set_cell_bg(self, cell, hex_color: str) -> None:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"),  "auto")
        shd.set(qn("w:fill"),   hex_color)
        tc_pr.append(shd)

    def _fmt_time(self, seconds: float) -> str:
        h = int(seconds // 3600)
        m = int((seconds % 3600) // 60)
        s = int(seconds % 60)
        return f"{h:02d}:{m:02d}:{s:02d}"
